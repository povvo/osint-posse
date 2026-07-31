#!/usr/bin/env python3
"""
ospo :: document intelligence

Local-only document metadata, text extraction, URL/email discovery, and
SHA-256 manifests. This module does not upload source files.

Extraction is best-effort. A successful run does not establish authorship,
authenticity, original creation time, or that all embedded content was parsed.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import mimetypes
import os
import re
import shutil
import stat
import subprocess
import sys
import tempfile
from collections.abc import Iterable, Iterator, Mapping, Sequence
from dataclasses import asdict, dataclass, field
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Literal
from urllib.parse import urlsplit, urlunsplit

from osint_common import ensure_dir, sha256_file, utc_now, write_json

try:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from docx import Document
    from docx.document import Document as DocumentType

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from PIL import ExifTags, Image, UnidentifiedImageError

    HAS_PIL = True
except ImportError:
    HAS_PIL = False


JsonDict = dict[str, Any]
ExtractionStatus = Literal["extracted", "not_applicable", "unavailable", "failed"]

MAX_FILE_BYTES = 512 * 1024 * 1024
MAX_TEXT_CHARS = 2_000_000
MAX_PDF_PAGES = 5_000
MAX_EXIFTOOL_OUTPUT_BYTES = 2_000_000
EXIFTOOL_TIMEOUT_SECONDS = 30

TEXT_SUFFIXES = frozenset(
    {
        ".csv",
        ".json",
        ".log",
        ".md",
        ".rst",
        ".text",
        ".txt",
        ".tsv",
        ".xml",
        ".yaml",
        ".yml",
    }
)
URL_RE = re.compile(r"https?://[^\s)\]>'\"<]+", re.IGNORECASE)
EMAIL_RE = re.compile(
    r"(?<![\w.+-])[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,63}(?![\w.-])",
    re.IGNORECASE,
)


class DocumentIntelError(RuntimeError):
    """Raised when an input path or workspace is unsafe or invalid."""


@dataclass(frozen=True)
class ExtractionPolicy:
    """Explicit limits for untrusted local documents."""

    max_file_bytes: int = MAX_FILE_BYTES
    max_text_chars: int = MAX_TEXT_CHARS
    max_pdf_pages: int = MAX_PDF_PAGES
    exiftool_timeout_seconds: int = EXIFTOOL_TIMEOUT_SECONDS

    def __post_init__(self) -> None:
        if self.max_file_bytes <= 0:
            raise ValueError("max_file_bytes must be positive.")
        if self.max_text_chars <= 0:
            raise ValueError("max_text_chars must be positive.")
        if self.max_pdf_pages <= 0:
            raise ValueError("max_pdf_pages must be positive.")
        if self.exiftool_timeout_seconds <= 0:
            raise ValueError("exiftool_timeout_seconds must be positive.")


@dataclass
class ExtractedText:
    """Text extraction result with explicit completion state."""

    status: ExtractionStatus
    text: str = ""
    extractor: str = ""
    truncated: bool = False
    warnings: list[str] = field(default_factory=list)


@dataclass
class DocumentFinding:
    """Per-file local analysis result and source-artifact identity."""

    file: str
    filename: str
    mime_type: str
    signature: str
    bytes: int
    modified_at_utc: str
    sha256: str
    text_extraction_status: ExtractionStatus
    text_extractor: str = ""
    extracted_text_chars: int = 0
    text_truncated: bool = False
    embedded_urls: list[str] = field(default_factory=list)
    embedded_emails: list[str] = field(default_factory=list)
    metadata: JsonDict = field(default_factory=dict)
    warnings: list[str] = field(default_factory=list)
    output_artifacts: JsonDict = field(default_factory=dict)


def _is_link_or_reparse(path: Path) -> bool:
    """Return whether path is a link, Windows junction, or reparse point."""
    result = path.lstat()
    attributes = getattr(result, "st_file_attributes", 0)
    reparse_flag = getattr(stat, "FILE_ATTRIBUTE_REPARSE_POINT", 0x400)
    return stat.S_ISLNK(result.st_mode) or bool(attributes & reparse_flag)


def _is_within(path: Path, root: Path) -> bool:
    """Return whether resolved path is contained by resolved root."""
    try:
        path.relative_to(root)
        return True
    except ValueError:
        return False


def _safe_files(root: Path, *, exclude_root: Path | None = None) -> Iterator[Path]:
    """Yield regular files without traversing links, reparse points, or exclusions."""
    resolved_root = root.resolve(strict=True)

    if _is_link_or_reparse(resolved_root):
        raise DocumentIntelError(f"Refusing link or reparse-point root: {root}")

    resolved_exclusion = (
        exclude_root.resolve(strict=False) if exclude_root is not None else None
    )
    pending = [resolved_root]

    while pending:
        directory = pending.pop()

        with os.scandir(directory) as entries:
            for entry in entries:
                item = Path(entry.path)

                if _is_link_or_reparse(item):
                    raise DocumentIntelError(
                        f"Directory scan refuses symlink or reparse point: {item}"
                    )

                resolved_item = item.resolve(strict=True)
                if not _is_within(resolved_item, resolved_root):
                    raise DocumentIntelError(
                        f"Scanned path escapes directory root: {item}"
                    )

                if (
                    resolved_exclusion is not None
                    and _is_within(resolved_item, resolved_exclusion)
                ):
                    continue

                if entry.is_dir(follow_symlinks=False):
                    pending.append(resolved_item)
                elif entry.is_file(follow_symlinks=False):
                    yield resolved_item


def _canonical_url(match: str) -> str:
    """Trim common trailing prose punctuation while preserving URL semantics."""
    value = match.rstrip(".,;:!?")
    parsed = urlsplit(value)

    if parsed.scheme.casefold() not in {"http", "https"} or not parsed.netloc:
        return value

    return urlunsplit(
        (
            parsed.scheme.casefold(),
            parsed.netloc,
            parsed.path,
            parsed.query,
            "",
        )
    )


def _deduplicated_matches(pattern: re.Pattern[str], text: str) -> list[str]:
    """Return case-insensitively deduplicated matches in deterministic order."""
    values = {match.group(0) for match in pattern.finditer(text)}
    return sorted(values, key=str.casefold)


def _safe_json_value(value: Any) -> Any:
    """Convert common metadata values to JSON-safe bounded representations."""
    if value is None or isinstance(value, (bool, int, float, str)):
        return value
    if isinstance(value, bytes):
        return value.hex()
    if isinstance(value, Mapping):
        return {
            str(key): _safe_json_value(item)
            for key, item in value.items()
        }
    if isinstance(value, Sequence) and not isinstance(value, (str, bytes)):
        return [_safe_json_value(item) for item in value]
    return str(value)


def _atomic_write_text(path: Path, text: str) -> None:
    """Atomically write a UTF-8 output artifact."""
    path.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.NamedTemporaryFile(
        mode="w",
        encoding="utf-8",
        newline="\n",
        dir=path.parent,
        prefix=f".{path.name}.",
        suffix=".tmp",
        delete=False,
    ) as handle:
        temporary_path = Path(handle.name)
        handle.write(text)
        handle.flush()
        os.fsync(handle.fileno())

    try:
        os.replace(temporary_path, path)
    finally:
        temporary_path.unlink(missing_ok=True)


class DocumentIntel:
    """Analyze local files without network retrieval or external upload."""

    def __init__(
        self,
        output_dir: str | Path = "./document_intel",
        *,
        policy: ExtractionPolicy | None = None,
    ) -> None:
        self.output_dir = Path(ensure_dir(output_dir)).resolve(strict=True)
        if _is_link_or_reparse(self.output_dir):
            raise DocumentIntelError(
                f"Output directory must not be a symlink or reparse point: "
                f"{self.output_dir}"
            )
        self.policy = policy or ExtractionPolicy()

    @staticmethod
    def _source_artifact_name(path: Path, digest: str) -> str:
        """Build an output-safe unique artifact stem from name and source hash."""
        safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", path.stem).strip("._")
        safe_name = safe_name[:80] or "document"
        return f"{safe_name}_{digest[:16]}"

    @staticmethod
    def _signature(path: Path) -> str:
        """Classify common magic bytes without treating it as authoritative MIME."""
        with path.open("rb") as handle:
            prefix = handle.read(16)

        if prefix.startswith(b"%PDF-"):
            return "pdf"
        if prefix.startswith(b"PK\x03\x04"):
            return "zip_container"
        if prefix.startswith(b"\x89PNG\r\n\x1a\n"):
            return "png"
        if prefix.startswith(b"\xff\xd8\xff"):
            return "jpeg"
        if prefix.startswith((b"GIF87a", b"GIF89a")):
            return "gif"
        if prefix.startswith(b"II*\x00") or prefix.startswith(b"MM\x00*"):
            return "tiff"
        return "unknown"

    def _bounded_text(self, text: str, *, extractor: str) -> ExtractedText:
        """Bound text output and visibly record truncation."""
        if len(text) <= self.policy.max_text_chars:
            return ExtractedText(
                status="extracted",
                text=text,
                extractor=extractor,
            )

        return ExtractedText(
            status="extracted",
            text=text[: self.policy.max_text_chars],
            extractor=extractor,
            truncated=True,
            warnings=[
                f"Extracted text truncated at {self.policy.max_text_chars} characters."
            ],
        )

    def _extract_pdf(self, path: Path) -> tuple[ExtractedText, JsonDict]:
        """Extract PDF text, metadata, and URI annotations without OCR."""
        if not HAS_PYPDF:
            return (
                ExtractedText(
                    status="unavailable",
                    extractor="pypdf",
                    warnings=["pypdf is not installed; PDF text was not extracted."],
                ),
                {},
            )

        try:
            reader = PdfReader(path, strict=False)
        except (OSError, PdfReadError, ValueError) as exc:
            return (
                ExtractedText(
                    status="failed",
                    extractor="pypdf",
                    warnings=[f"PDF open failed: {exc}"],
                ),
                {},
            )

        metadata: JsonDict = {
            "pdf": {
                "page_count": len(reader.pages),
                "is_encrypted": reader.is_encrypted,
                "document_metadata": {
                    str(key): str(value)
                    for key, value in (reader.metadata or {}).items()
                },
                "annotation_urls": [],
            }
        }

        if reader.is_encrypted:
            try:
                decrypt_result = reader.decrypt("")
            except (PdfReadError, ValueError) as exc:
                return (
                    ExtractedText(
                        status="failed",
                        extractor="pypdf",
                        warnings=[f"Encrypted PDF could not be opened: {exc}"],
                    ),
                    metadata,
                )
            if decrypt_result == 0:
                return (
                    ExtractedText(
                        status="unavailable",
                        extractor="pypdf",
                        warnings=[
                            "PDF is encrypted and could not be opened with an empty password."
                        ],
                    ),
                    metadata,
                )

        page_count = len(reader.pages)
        inspected_pages = min(page_count, self.policy.max_pdf_pages)
        text_parts: list[str] = []
        annotation_urls: set[str] = set()
        warnings: list[str] = []

        for page_number in range(inspected_pages):
            page = reader.pages[page_number]
            try:
                text_parts.append(page.extract_text() or "")
            except (PdfReadError, ValueError) as exc:
                warnings.append(f"PDF page {page_number + 1} text extraction failed: {exc}")

            annotations = page.get("/Annots", [])
            for annotation_ref in annotations:
                try:
                    annotation = annotation_ref.get_object()
                    action = annotation.get("/A", {})
                    uri = action.get("/URI") if isinstance(action, Mapping) else None
                    if uri:
                        annotation_urls.add(str(uri))
                except (AttributeError, KeyError, PdfReadError, ValueError):
                    warnings.append(
                        f"PDF page {page_number + 1} contains unreadable annotation data."
                    )

        if page_count > inspected_pages:
            warnings.append(
                f"Only the first {inspected_pages} of {page_count} PDF pages were inspected."
            )

        metadata["pdf"]["annotation_urls"] = sorted(annotation_urls, key=str.casefold)
        extracted = self._bounded_text("\n".join(text_parts), extractor="pypdf")
        extracted.warnings.extend(warnings)

        if not extracted.text and not extracted.warnings:
            extracted.warnings.append(
                "PDF contained no extractable text; it may be image-only or use unsupported encoding."
            )
        return extracted, metadata

    @staticmethod
    def _paragraph_text(paragraphs: Iterable[Any]) -> list[str]:
        """Collect non-empty paragraph text from python-docx structures."""
        return [
            paragraph.text
            for paragraph in paragraphs
            if isinstance(paragraph.text, str) and paragraph.text.strip()
        ]

    def _extract_docx(self, path: Path) -> tuple[ExtractedText, JsonDict]:
        """Extract DOCX body/table/header/footer text and core properties."""
        if not HAS_DOCX:
            return (
                ExtractedText(
                    status="unavailable",
                    extractor="python-docx",
                    warnings=[
                        "python-docx is not installed; DOCX text was not extracted."
                    ],
                ),
                {},
            )

        try:
            document: DocumentType = Document(path)
        except (OSError, ValueError, KeyError) as exc:
            return (
                ExtractedText(
                    status="failed",
                    extractor="python-docx",
                    warnings=[f"DOCX open failed: {exc}"],
                ),
                {},
            )

        lines = self._paragraph_text(document.paragraphs)
        table_cells = 0
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    lines.append(" | ".join(values))
                    table_cells += len(values)

        header_footer_lines = 0
        for section in document.sections:
            for container in (
                section.header,
                section.footer,
                section.first_page_header,
                section.first_page_footer,
                section.even_page_header,
                section.even_page_footer,
            ):
                extracted_lines = self._paragraph_text(container.paragraphs)
                lines.extend(extracted_lines)
                header_footer_lines += len(extracted_lines)

        core = document.core_properties
        metadata: JsonDict = {
            "docx": {
                "core_properties": {
                    key: _safe_json_value(getattr(core, key, None))
                    for key in (
                        "author",
                        "category",
                        "comments",
                        "content_status",
                        "created",
                        "identifier",
                        "keywords",
                        "language",
                        "last_modified_by",
                        "last_printed",
                        "modified",
                        "revision",
                        "subject",
                        "title",
                        "version",
                    )
                },
                "table_cell_values_extracted": table_cells,
                "header_footer_paragraphs_extracted": header_footer_lines,
            }
        }

        return self._bounded_text("\n".join(lines), extractor="python-docx"), metadata

    def _extract_text_file(self, path: Path) -> tuple[ExtractedText, JsonDict]:
        """Read UTF-8 text in bounded chunks, preserving replacement behavior."""
        chunks: list[str] = []
        characters_read = 0
        truncated = False

        try:
            with path.open("r", encoding="utf-8", errors="replace") as handle:
                while characters_read < self.policy.max_text_chars:
                    remaining = self.policy.max_text_chars - characters_read
                    chunk = handle.read(min(64 * 1024, remaining))
                    if not chunk:
                        break
                    chunks.append(chunk)
                    characters_read += len(chunk)

                if handle.read(1):
                    truncated = True
        except OSError as exc:
            return (
                ExtractedText(
                    status="failed",
                    extractor="utf-8",
                    warnings=[f"Text read failed: {exc}"],
                ),
                {},
            )

        warnings = (
            [f"Text content truncated at {self.policy.max_text_chars} characters."]
            if truncated
            else []
        )
        return (
            ExtractedText(
                status="extracted",
                text="".join(chunks),
                extractor="utf-8",
                truncated=truncated,
                warnings=warnings,
            ),
            {},
        )

    def _extract_image_metadata(self, path: Path) -> tuple[ExtractedText, JsonDict]:
        """Extract basic image metadata and selected EXIF fields without OCR."""
        if not HAS_PIL:
            return (
                ExtractedText(
                    status="unavailable",
                    extractor="Pillow",
                    warnings=[
                        "Pillow is not installed; image metadata was not extracted."
                    ],
                ),
                {},
            )

        try:
            with Image.open(path) as image:
                exif: JsonDict = {}
                for tag_id, value in image.getexif().items():
                    tag_name = ExifTags.TAGS.get(tag_id, str(tag_id))
                    exif[str(tag_name)] = _safe_json_value(value)

                metadata: JsonDict = {
                    "image": {
                        "format": image.format,
                        "mode": image.mode,
                        "width": image.width,
                        "height": image.height,
                        "frames": getattr(image, "n_frames", 1),
                        "exif": exif,
                    }
                }
        except (OSError, UnidentifiedImageError, ValueError) as exc:
            return (
                ExtractedText(
                    status="failed",
                    extractor="Pillow",
                    warnings=[f"Image metadata extraction failed: {exc}"],
                ),
                {},
            )

        return (
            ExtractedText(
                status="not_applicable",
                extractor="Pillow",
                warnings=["OCR is not performed; no image text extraction attempted."],
            ),
            metadata,
        )

    def _extract_text_and_metadata(
        self,
        path: Path,
        mime_type: str,
    ) -> tuple[ExtractedText, JsonDict]:
        """Select a bounded local extractor based on suffix and MIME type."""
        suffix = path.suffix.casefold()

        if suffix == ".pdf":
            return self._extract_pdf(path)
        if suffix == ".docx":
            return self._extract_docx(path)
        if mime_type.startswith("text/") or suffix in TEXT_SUFFIXES:
            return self._extract_text_file(path)
        if mime_type.startswith("image/"):
            return self._extract_image_metadata(path)

        return (
            ExtractedText(
                status="not_applicable",
                warnings=["No local extractor is configured for this file type."],
            ),
            {},
        )

    def _exiftool_metadata(self, path: Path) -> tuple[JsonDict, list[str]]:
        """Read ExifTool JSON metadata if available, without suppressing failures."""
        if shutil.which("exiftool") is None:
            return {}, ["ExifTool not installed; extended metadata was not collected."]

        try:
            completed = subprocess.run(
                ["exiftool", "-j", "--", str(path)],
                capture_output=True,
                check=False,
                timeout=self.policy.exiftool_timeout_seconds,
            )
        except subprocess.TimeoutExpired:
            return {}, [
                f"ExifTool exceeded {self.policy.exiftool_timeout_seconds} seconds."
            ]
        except OSError as exc:
            return {}, [f"ExifTool could not be executed: {exc}"]

        if completed.returncode != 0:
            error = completed.stderr.decode("utf-8", errors="replace")[:1000]
            return {}, [f"ExifTool failed with exit {completed.returncode}: {error}"]

        if len(completed.stdout) > MAX_EXIFTOOL_OUTPUT_BYTES:
            return {}, [
                "ExifTool metadata output exceeded the configured safety limit."
            ]

        try:
            value = json.loads(completed.stdout)
        except json.JSONDecodeError as exc:
            return {}, [f"ExifTool returned invalid JSON: {exc}"]

        if not isinstance(value, list) or not value or not isinstance(value[0], dict):
            return {}, ["ExifTool returned an unexpected JSON structure."]

        return _safe_json_value(value[0]), []

    def analyse_file(
        self,
        path: str | Path,
        *,
        save_text: bool = True,
    ) -> JsonDict:
        """Analyze one regular local file and write its separate result artifact."""
        source_path = Path(path).expanduser()

        if not source_path.exists():
            raise FileNotFoundError(f"File not found: {source_path}")
        if _is_link_or_reparse(source_path):
            raise DocumentIntelError(
                f"Refusing symlink or reparse-point input file: {source_path}"
            )
        if not source_path.is_file():
            raise DocumentIntelError(f"Input is not a regular file: {source_path}")

        source_path = source_path.resolve(strict=True)
        file_size = source_path.stat().st_size
        if file_size > self.policy.max_file_bytes:
            raise DocumentIntelError(
                f"File exceeds {self.policy.max_file_bytes} byte analysis limit: "
                f"{source_path}"
            )

        digest = sha256_file(source_path)
        mime_type, _encoding = mimetypes.guess_type(source_path.name)
        mime_type = mime_type or "application/octet-stream"
        signature = self._signature(source_path)

        extracted, extractor_metadata = self._extract_text_and_metadata(
            source_path,
            mime_type,
        )
        exiftool_metadata, exiftool_warnings = self._exiftool_metadata(source_path)

        all_urls = {
            _canonical_url(value)
            for value in URL_RE.findall(extracted.text)
        }
        pdf_urls = extractor_metadata.get("pdf", {}).get("annotation_urls", [])
        if isinstance(pdf_urls, list):
            all_urls.update(str(value) for value in pdf_urls)

        metadata = dict(extractor_metadata)
        if exiftool_metadata:
            metadata["exiftool"] = exiftool_metadata

        artifact_stem = self._source_artifact_name(source_path, digest)
        text_artifact: str | None = None
        if save_text and extracted.text:
            text_path = self.output_dir / f"{artifact_stem}.extracted.txt"
            _atomic_write_text(text_path, extracted.text)
            text_artifact = str(text_path)

        finding = DocumentFinding(
            file=str(source_path),
            filename=source_path.name,
            mime_type=mime_type,
            signature=signature,
            bytes=file_size,
            modified_at_utc=datetime.fromtimestamp(
                source_path.stat().st_mtime,
                UTC,
            ).isoformat(timespec="seconds"),
            sha256=digest,
            text_extraction_status=extracted.status,
            text_extractor=extracted.extractor,
            extracted_text_chars=len(extracted.text),
            text_truncated=extracted.truncated,
            embedded_urls=sorted(all_urls, key=str.casefold),
            embedded_emails=_deduplicated_matches(EMAIL_RE, extracted.text),
            metadata=metadata,
            warnings=extracted.warnings + exiftool_warnings,
            output_artifacts={
                "extracted_text": text_artifact,
            },
        )

        result: JsonDict = {
            "schema_version": "1.0",
            "analysed_at_utc": utc_now(),
            "scope_notice": (
                "Local analysis only; no source file content was uploaded by this tool."
            ),
            "finding": asdict(finding),
        }
        report_path = self.output_dir / f"{artifact_stem}.document_intel.json"
        write_json(report_path, result)
        result["output_path"] = str(report_path)
        return result

    def analyse_directory(self, directory: str | Path) -> JsonDict:
        """Analyze all safe regular files below one directory and write a manifest."""
        root = Path(directory).expanduser()
        if not root.exists():
            raise FileNotFoundError(f"Directory not found: {root}")
        if _is_link_or_reparse(root):
            raise DocumentIntelError(
                f"Refusing symlink or reparse-point directory: {root}"
            )
        if not root.is_dir():
            raise NotADirectoryError(f"Input is not a directory: {root}")

        root = root.resolve(strict=True)
        if _is_within(self.output_dir, root):
            raise DocumentIntelError(
                "Output directory is inside the scanned directory. Choose an output "
                "directory outside the scan root to prevent self-analysis."
            )

        files: list[JsonDict] = []
        failures: list[JsonDict] = []

        for item in sorted(_safe_files(root), key=lambda value: value.as_posix()):
            try:
                result = self.analyse_file(item, save_text=False)
            except (DocumentIntelError, FileNotFoundError, OSError) as exc:
                failures.append(
                    {
                        "file": str(item.relative_to(root)),
                        "error": str(exc),
                    }
                )
                continue

            finding = result["finding"]
            finding["file"] = str(item.relative_to(root))
            files.append(finding)

        summary: JsonDict = {
            "schema_version": "1.0",
            "directory": str(root),
            "analysed_at_utc": utc_now(),
            "scope_notice": (
                "Local analysis only; no source file content was uploaded by this tool."
            ),
            "policy": asdict(self.policy),
            "files": files,
            "failures": failures,
            "counts": {
                "total_files": len(files) + len(failures),
                "analysed_files": len(files),
                "failed_files": len(failures),
                "total_bytes": sum(int(finding["bytes"]) for finding in files),
            },
        }
        manifest_path = self.output_dir / "directory_document_intel.json"
        write_json(manifest_path, summary)
        summary["output_path"] = str(manifest_path)
        return summary


def main() -> int:
    """Run local document intelligence from the command line."""
    parser = argparse.ArgumentParser(
        description="Extract local document metadata and text without uploading files"
    )
    parser.add_argument("path", type=Path)
    parser.add_argument("--out-dir", type=Path, default=Path("./document_intel"))
    parser.add_argument(
        "--directory",
        action="store_true",
        help="Recursively analyze regular files below path",
    )
    parser.add_argument(
        "--max-file-mb",
        type=int,
        default=MAX_FILE_BYTES // (1024 * 1024),
        help="Maximum source file size to analyze",
    )
    parser.add_argument(
        "--max-text-chars",
        type=int,
        default=MAX_TEXT_CHARS,
        help="Maximum extracted characters retained per file",
    )
    args = parser.parse_args()

    try:
        policy = ExtractionPolicy(
            max_file_bytes=args.max_file_mb * 1024 * 1024,
            max_text_chars=args.max_text_chars,
        )
        intel = DocumentIntel(args.out_dir, policy=policy)
        result = (
            intel.analyse_directory(args.path)
            if args.directory
            else intel.analyse_file(args.path)
        )
    except (DocumentIntelError, FileNotFoundError, NotADirectoryError, ValueError) as exc:
        print(f"Document analysis failed: {exc}", file=sys.stderr)
        return 2

    print(json.dumps(result, indent=2, sort_keys=True))
    return 1 if result.get("counts", {}).get("failed_files", 0) else 0


if __name__ == "__main__":
    raise SystemExit(main())